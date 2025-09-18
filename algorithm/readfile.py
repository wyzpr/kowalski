from docx import Document
from pypdf import PdfReader
from langchain.tools import Tool
import os

def readDoc(docx_path):
    """
    Extracts text from a Word document (.docx) including headers, footers, paragraphs
    and tables.
    """
    doc = Document(docx_path)
    fullText = []
    for section in doc.sections:
        print("Section found")
        header = section.header
        if header is not None:
            for paragraph in header.paragraphs:
                fullText.append(paragraph.text)
        footer = section.footer
        if footer is not None:
            for paragraph in footer.paragraphs:
                fullText.append(paragraph.text)
    
    for paragraph in doc.paragraphs:
        print("Paragraph found")
        fullText.append(paragraph.text)

    for table in doc.tables:
        print("Table found")
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            fullText.append(' | '.join(row_text))
    
    return '\n'.join(fullText)


def readPDF(pdf_path):
    """
    Extracts text from a PDF file using pypdf.
    """
    reader = PdfReader(pdf_path)
    full_text = []
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        full_text.append(page.extract_text())
    return '\n'.join(full_text)


def readFile(filePath):
    if os.path.exists(filePath) == False:
        print("File does not exist. Please provide a valid file path.")
        exit()

    if '.docx' in filePath:
        content = readDoc(filePath)
    elif '.pdf' in filePath:
        content = readPDF(filePath)
    else:
        print("Unsupported file format. Please provide a .docx or .pdf file.")
        exit()
    return content

read_file = Tool(
    name="read_file",
    func=readFile,
    description="Returns the text content of a .docx or .pdf file given its file path. Use this tool to extract text from CV files.",
)